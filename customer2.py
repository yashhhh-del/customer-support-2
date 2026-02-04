import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import uuid
import time
import json
import sqlite3
from typing import List, Dict, Optional, Tuple
import re
from io import BytesIO
import base64

# For PDF/DOCX/Excel processing
try:
    import PyPDF2
    from docx import Document
    import openpyxl
    from openpyxl import load_workbook
    DOCUMENT_SUPPORT = True
except ImportError:
    DOCUMENT_SUPPORT = False
    st.warning("⚠️ Install document support: pip install PyPDF2 python-docx openpyxl")

# For email functionality
try:
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    import imaplib
    import email
    EMAIL_SUPPORT = True
except ImportError:
    EMAIL_SUPPORT = False

# For WhatsApp integration
import urllib.parse

# For vector database and embeddings (optional)
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
    from langchain.chains import RetrievalQA
    from langchain.docstore.document import Document as LangchainDocument
    from langchain.prompts import PromptTemplate
    AI_SUPPORT = True
except ImportError:
    AI_SUPPORT = False

# For language detection and translation (optional)
try:
    from langdetect import detect, LangDetectException
    from deep_translator import GoogleTranslator
    TRANSLATION_SUPPORT = True
except ImportError:
    TRANSLATION_SUPPORT = False

# For web scraping (optional)
try:
    import requests
    from bs4 import BeautifulSoup
    WEB_SCRAPE_SUPPORT = True
except ImportError:
    WEB_SCRAPE_SUPPORT = False

# For OCR (optional)
try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False

# For semantic similarity (optional)
try:
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    SIMILARITY_SUPPORT = True
except ImportError:
    SIMILARITY_SUPPORT = False

# Page configuration
st.set_page_config(
    page_title="AI Support Agent Pro",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced Custom CSS
st.markdown("""
<style>
    :root {
        --primary-blue: #0066CC;
        --primary-dark: #003D7A;
        --success-green: #00A86B;
        --warning-orange: #FF8C00;
        --danger-red: #DC143C;
        --bg-light: #F8F9FA;
        --bg-white: #FFFFFF;
        --text-dark: #212529;
        --text-light: #6C757D;
        --border-color: #DEE2E6;
    }
    
    .main-header {
        font-size: 2.8rem;
        font-weight: 800;
        background: linear-gradient(135deg, #0066CC 0%, #00A86B 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
    }
    
    .sub-header {
        font-size: 1.5rem;
        font-weight: 600;
        color: var(--primary-dark);
        margin-bottom: 1rem;
        border-bottom: 3px solid var(--primary-blue);
        padding-bottom: 0.5rem;
    }
    
    .chat-message {
        padding: 1.2rem;
        border-radius: 12px;
        margin-bottom: 1rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border: 1px solid var(--border-color);
    }
    
    .user-message {
        background: linear-gradient(135deg, #E3F2FD 0%, #BBDEFB 100%);
        margin-left: 3rem;
        border-left: 4px solid #0066CC;
        color: #003D7A;
    }
    
    .bot-message {
        background: linear-gradient(135deg, #F5F5F5 0%, #E8E8E8 100%);
        margin-right: 3rem;
        border-left: 4px solid #00A86B;
        color: #212529;
    }
    
    .confidence-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
        margin: 0 4px;
    }
    
    .confidence-high {
        background-color: #00A86B;
        color: #FFFFFF;
    }
    
    .confidence-medium {
        background-color: #FF8C00;
        color: #FFFFFF;
    }
    
    .confidence-low {
        background-color: #DC143C;
        color: #FFFFFF;
    }
    
    .ticket-card {
        padding: 1.5rem;
        border: 2px solid var(--border-color);
        border-radius: 12px;
        margin-bottom: 1rem;
        background: var(--bg-white);
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
    }
    
    .ticket-card:hover {
        box-shadow: 0 6px 20px rgba(0,0,0,0.15);
        transform: translateY(-2px);
    }
    
    .priority-high {
        background-color: #DC143C;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    .priority-medium {
        background-color: #FF8C00;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    .priority-low {
        background-color: #00A86B;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    .status-open {
        background-color: #0066CC;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    .status-progress {
        background-color: #FF8C00;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    .status-closed {
        background-color: #6C757D;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    .contact-banner {
        background: linear-gradient(135deg, #0066CC 0%, #00A86B 100%);
        padding: 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 8px 24px rgba(0,0,0,0.15);
    }
    
    .contact-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        min-width: 280px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    
    .contact-card:hover {
        transform: translateY(-8px);
        box-shadow: 0 12px 28px rgba(0,0,0,0.25);
        background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
    }
    
    .contact-card h4 {
        color: var(--primary-dark);
        margin-top: 0;
        font-weight: 700;
    }
    
    .contact-card p {
        color: var(--text-dark);
        font-size: 1rem;
        font-weight: 500;
        margin: 10px 0;
    }
    
    .whatsapp-button {
        background-color: #25D366;
        color: white;
        padding: 12px 24px;
        border-radius: 8px;
        text-decoration: none;
        display: inline-block;
        margin: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        border: 2px solid #128C7E;
    }
    
    .whatsapp-button:hover {
        background-color: #128C7E;
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(37, 211, 102, 0.4);
    }
    
    .email-button {
        background-color: #0066CC;
        color: white;
        padding: 12px 24px;
        border-radius: 8px;
        text-decoration: none;
        display: inline-block;
        margin: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        border: 2px solid #003D7A;
    }
    
    .email-button:hover {
        background-color: #003D7A;
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(0, 102, 204, 0.4);
    }
    
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        border: 2px solid var(--border-color);
        text-align: center;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 800;
        color: var(--primary-blue);
    }
    
    .metric-label {
        font-size: 1rem;
        color: var(--text-light);
        font-weight: 600;
        margin-top: 0.5rem;
    }
    
    .info-box {
        background: #E3F2FD;
        border-left: 5px solid #0066CC;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #003D7A;
        font-weight: 500;
    }
    
    .success-box {
        background: #D4EDDA;
        border-left: 5px solid #00A86B;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #155724;
        font-weight: 500;
    }
    
    .warning-box {
        background: #FFF3CD;
        border-left: 5px solid #FF8C00;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #856404;
        font-weight: 500;
    }
    
    .error-box {
        background: #F8D7DA;
        border-left: 5px solid #DC143C;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #721C24;
        font-weight: 500;
    }
    
    .demo-mode-banner {
        background: linear-gradient(135deg, #FF8C00 0%, #DC143C 100%);
        color: white;
        padding: 1rem;
        border-radius: 8px;
        text-align: center;
        font-weight: 600;
        margin-bottom: 1rem;
    }
    
    .footer {
        text-align: center;
        padding: 2rem;
        background: var(--bg-light);
        border-top: 3px solid var(--primary-blue);
        margin-top: 3rem;
        color: var(--text-dark);
    }
</style>
""", unsafe_allow_html=True)

# Database setup
def init_database():
    """Initialize SQLite database for persistence"""
    conn = sqlite3.connect('support_agent.db', check_same_thread=False)
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS chat_history
                 (id TEXT PRIMARY KEY, role TEXT, message TEXT, timestamp TEXT, 
                  confidence REAL, response_time REAL, language TEXT, category TEXT)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tickets
                 (id TEXT PRIMARY KEY, query TEXT, language TEXT, category TEXT, 
                  status TEXT, priority TEXT, assigned_to TEXT, timestamp TEXT, 
                  resolved_at TEXT, resolution_time REAL, channel TEXT)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS feedback
                 (chat_id TEXT, feedback TEXT, timestamp TEXT, comment TEXT)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS analytics
                 (date TEXT, total_queries INTEGER, answered INTEGER, escalated INTEGER,
                  avg_response_time REAL, avg_confidence REAL)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS faq_data
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, question TEXT, answer TEXT, 
                  category TEXT, language TEXT)''')
    
    conn.commit()
    return conn

# Initialize database
if 'db_conn' not in st.session_state:
    st.session_state.db_conn = init_database()

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'tickets' not in st.session_state:
    st.session_state.tickets = []
if 'analytics' not in st.session_state:
    st.session_state.analytics = {
        'total_queries': 0,
        'answered': 0,
        'escalated': 0,
        'languages': {'English': 0, 'Hindi': 0, 'Marathi': 0, 'Other': 0},
        'categories': {'Billing': 0, 'Technical': 0, 'General': 0, 'Product': 0}
    }
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'knowledge_base_text' not in st.session_state:
    st.session_state.knowledge_base_text = ""
if 'faq_database' not in st.session_state:
    st.session_state.faq_database = []
if 'feedback' not in st.session_state:
    st.session_state.feedback = {}
if 'embeddings_model' not in st.session_state:
    st.session_state.embeddings_model = None
if 'channel_stats' not in st.session_state:
    st.session_state.channel_stats = {
        'Website': 0,
        'WhatsApp': 0,
        'Email': 0,
        'Manual': 0
    }
if 'email_queue' not in st.session_state:
    st.session_state.email_queue = []
if 'whatsapp_messages' not in st.session_state:
    st.session_state.whatsapp_messages = []
if 'gmail_config' not in st.session_state:
    st.session_state.gmail_config = {
        'email': 'support@yourcompany.com',
        'smtp_server': 'smtp.gmail.com',
        'smtp_port': 587,
        'imap_server': 'imap.gmail.com'
    }
if 'whatsapp_config' not in st.session_state:
    st.session_state.whatsapp_config = {
        'phone_number': '+1234567890',
        'wa_link': 'https://wa.me/1234567890'
    }
if 'kb_processed' not in st.session_state:
    st.session_state.kb_processed = False
if 'demo_mode' not in st.session_state:
    st.session_state.demo_mode = True

# Helper Functions
def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    if not DOCUMENT_SUPPORT:
        return ""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"❌ Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    if not DOCUMENT_SUPPORT:
        return ""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"❌ Error reading DOCX: {str(e)}")
        return ""

def extract_faq_from_excel(file) -> List[Dict]:
    """Extract FAQ data from Excel file"""
    if not DOCUMENT_SUPPORT:
        return []
    try:
        df = pd.read_excel(file)
        
        # Try to identify question and answer columns
        possible_question_cols = ['question', 'query', 'q', 'questions', 'प्रश्न', 'सवाल']
        possible_answer_cols = ['answer', 'response', 'a', 'answers', 'उत्तर', 'जवाब']
        possible_category_cols = ['category', 'type', 'श्रेणी']
        
        question_col = None
        answer_col = None
        category_col = None
        
        # Find question column
        for col in df.columns:
            if col.lower() in possible_question_cols:
                question_col = col
                break
        
        # Find answer column
        for col in df.columns:
            if col.lower() in possible_answer_cols:
                answer_col = col
                break
        
        # Find category column
        for col in df.columns:
            if col.lower() in possible_category_cols:
                category_col = col
                break
        
        # If columns not found, use first two columns
        if question_col is None and len(df.columns) >= 1:
            question_col = df.columns[0]
        if answer_col is None and len(df.columns) >= 2:
            answer_col = df.columns[1]
        
        faq_list = []
        for _, row in df.iterrows():
            if pd.notna(row.get(question_col)) and pd.notna(row.get(answer_col)):
                faq_list.append({
                    'question': str(row[question_col]),
                    'answer': str(row[answer_col]),
                    'category': str(row[category_col]) if category_col and pd.notna(row.get(category_col)) else 'General'
                })
        
        return faq_list
    except Exception as e:
        st.error(f"❌ Error reading Excel FAQ: {str(e)}")
        return []

def extract_text_from_excel(file) -> str:
    """Extract text from Excel file (XLSX/XLS)"""
    if not DOCUMENT_SUPPORT:
        return ""
    try:
        wb = load_workbook(file, read_only=True, data_only=True)
        text = ""
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"\n[Sheet: {sheet_name}]\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) for cell in row if cell is not None]
                if row_data:
                    text += " | ".join(row_data) + "\n"
        
        wb.close()
        return text
    except Exception as e:
        st.error(f"❌ Error reading Excel file: {str(e)}")
        return ""

def detect_language(text: str) -> str:
    """Detect language of text"""
    if not TRANSLATION_SUPPORT:
        return 'English'
    try:
        lang_code = detect(text)
        lang_map = {
            'en': 'English',
            'hi': 'Hindi',
            'mr': 'Marathi',
        }
        return lang_map.get(lang_code, 'Other')
    except:
        return 'English'

def translate_text(text: str, target_lang: str) -> str:
    """Translate text to target language"""
    if not TRANSLATION_SUPPORT or target_lang == 'English':
        return text
    try:
        lang_code_map = {
            'Hindi': 'hi',
            'Marathi': 'mr',
        }
        target_code = lang_code_map.get(target_lang, 'en')
        translator = GoogleTranslator(source='auto', target=target_code)
        translated = translator.translate(text)
        return translated
    except Exception as e:
        return text

def categorize_query(query: str) -> str:
    """Enhanced keyword-based categorization"""
    query_lower = query.lower()
    
    billing_keywords = ['payment', 'invoice', 'bill', 'charge', 'refund', 'price', 'cost', 
                       'subscription', 'card', 'billing', 'money', 'paid', 'transaction',
                       'भुगतान', 'पैसा', 'बिल']
    technical_keywords = ['error', 'bug', 'issue', 'problem', 'not working', 'broken', 
                         'crash', 'slow', 'loading', 'login', 'access', 'technical', 'password',
                         'installation', 'update', 'sync', 'त्रुटि', 'समस्या']
    product_keywords = ['feature', 'how to', 'tutorial', 'guide', 'demo', 'product', 
                       'functionality', 'use', 'work', 'setup', 'configure', 'कैसे']
    
    billing_score = sum(1 for keyword in billing_keywords if keyword in query_lower)
    technical_score = sum(1 for keyword in technical_keywords if keyword in query_lower)
    product_score = sum(1 for keyword in product_keywords if keyword in query_lower)
    
    scores = {
        'Billing': billing_score,
        'Technical': technical_score,
        'Product': product_score
    }
    
    max_category = max(scores, key=scores.get)
    return max_category if scores[max_category] > 0 else 'General'

def assign_priority(confidence: float, category: str) -> str:
    """Assign priority based on confidence and category"""
    if confidence < 0.4 or category == 'Billing':
        return 'High'
    elif confidence < 0.6 or category == 'Technical':
        return 'Medium'
    else:
        return 'Low'

def fuzzy_match_faq(query: str, faq_list: List[Dict], threshold: float = 0.6) -> Tuple[Optional[str], float]:
    """Simple fuzzy matching for FAQ (demo mode)"""
    query_lower = query.lower()
    query_words = set(query_lower.split())
    
    best_match = None
    best_score = 0
    
    for faq in faq_list:
        question_lower = faq['question'].lower()
        question_words = set(question_lower.split())
        
        # Calculate word overlap
        common_words = query_words.intersection(question_words)
        if len(query_words) > 0:
            score = len(common_words) / len(query_words)
        else:
            score = 0
        
        # Check for exact substring match
        if query_lower in question_lower or question_lower in query_lower:
            score = max(score, 0.8)
        
        if score > best_score:
            best_score = score
            best_match = faq['answer']
    
    return best_match, best_score

def get_demo_response(query: str, faq_list: List[Dict]) -> Tuple[str, float]:
    """Get response in demo mode using FAQ matching"""
    if faq_list:
        answer, confidence = fuzzy_match_faq(query, faq_list)
        if answer and confidence > 0.3:
            return answer, confidence
    
    # Default response
    default_response = """Thank you for your question! 

I'm currently running in demo mode. To get better AI-powered responses, please:
1. Enter your OpenAI API key in the sidebar
2. Click "Process Knowledge Base"

For immediate assistance, please contact our support team via:
📧 Email or 💬 WhatsApp (see contact info above)

Your query has been noted and will be reviewed by our team."""
    
    return default_response, 0.5

def get_ai_response(query: str, vector_store, openai_api_key: str, target_language: str = 'English') -> Tuple[str, float, List]:
    """Get AI response using RAG with multilingual support"""
    if not AI_SUPPORT or not vector_store:
        # Fallback to demo mode
        answer, confidence = get_demo_response(query, st.session_state.faq_database)
        return answer, confidence, []
    
    try:
        english_query = query
        if target_language != 'English' and TRANSLATION_SUPPORT:
            try:
                translator = GoogleTranslator(source='auto', target='en')
                english_query = translator.translate(query)
            except:
                pass
        
        llm = ChatOpenAI(
            model_name="gpt-3.5-turbo",
            temperature=0.3,
            openai_api_key=openai_api_key
        )
        
        prompt_template = """You are a helpful and professional customer support assistant. 
        Use the following context to answer the question accurately and concisely. 
        If you cannot find the answer in the context, politely inform the customer and suggest contacting support for further assistance.
        
        Context: {context}
        
        Question: {question}
        
        Provide a helpful, accurate, and friendly answer:"""
        
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 4}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        result = qa_chain({"query": english_query})
        
        answer = result['result']
        source_docs = result['source_documents']
        
        # Calculate confidence
        confidence = 0.75 if len(source_docs) > 0 else 0.5
        
        if target_language != 'English' and TRANSLATION_SUPPORT:
            answer = translate_text(answer, target_language)
        
        return answer, confidence, source_docs
    except Exception as e:
        st.error(f"❌ Error getting AI response: {str(e)}")
        # Fallback to demo mode
        answer, confidence = get_demo_response(query, st.session_state.faq_database)
        return answer, confidence, []

def create_ticket(query: str, language: str, category: str, confidence: float, channel: str = 'Website'):
    """Create escalation ticket with priority and assignment"""
    priority = assign_priority(confidence, category)
    
    agents = ['Agent A', 'Agent B', 'Agent C', 'Agent D']
    assigned_to = agents[len(st.session_state.tickets) % len(agents)]
    
    ticket = {
        'id': str(uuid.uuid4())[:8].upper(),
        'query': query,
        'language': language,
        'category': category,
        'status': 'Open',
        'priority': priority,
        'assigned_to': assigned_to,
        'channel': channel,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'resolved_at': None,
        'resolution_time': None
    }
    
    st.session_state.tickets.append(ticket)
    st.session_state.analytics['escalated'] += 1
    
    conn = st.session_state.db_conn
    c = conn.cursor()
    c.execute('''INSERT INTO tickets VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
              (ticket['id'], ticket['query'], ticket['language'], ticket['category'],
               ticket['status'], ticket['priority'], ticket['assigned_to'], 
               ticket['timestamp'], ticket['resolved_at'], ticket['resolution_time'],
               ticket['channel']))
    conn.commit()
    
    return ticket['id']

def save_chat_to_db(chat_entry: Dict):
    """Save chat entry to database"""
    conn = st.session_state.db_conn
    c = conn.cursor()
    
    chat_id = str(uuid.uuid4())
    c.execute('''INSERT INTO chat_history VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
              (chat_id, chat_entry['role'], chat_entry['message'], 
               chat_entry['timestamp'].strftime("%Y-%m-%d %H:%M:%S"),
               chat_entry.get('confidence', None),
               chat_entry.get('response_time', None),
               chat_entry.get('language', None),
               chat_entry.get('category', None)))
    conn.commit()
    return chat_id

def create_whatsapp_link(phone_number: str, message: str) -> str:
    """Create WhatsApp direct message link"""
    encoded_message = urllib.parse.quote(message)
    phone_clean = phone_number.replace('+', '').replace('-', '').replace(' ', '')
    return f"https://wa.me/{phone_clean}?text={encoded_message}"

def get_confidence_badge(confidence: float) -> str:
    """Generate HTML badge for confidence score"""
    if confidence >= 0.7:
        return f'<span class="confidence-badge confidence-high">🟢 High: {confidence:.1%}</span>'
    elif confidence >= 0.5:
        return f'<span class="confidence-badge confidence-medium">🟡 Medium: {confidence:.1%}</span>'
    else:
        return f'<span class="confidence-badge confidence-low">🔴 Low: {confidence:.1%}</span>'

def get_priority_badge(priority: str) -> str:
    """Generate HTML badge for priority"""
    return f'<span class="priority-{priority.lower()}">{priority}</span>'

def get_status_badge(status: str) -> str:
    """Generate HTML badge for status"""
    status_map = {
        'Open': 'open',
        'In Progress': 'progress',
        'Closed': 'closed'
    }
    return f'<span class="status-{status_map.get(status, "open")}">{status}</span>'

# Sidebar - Knowledge Management
with st.sidebar:
    st.markdown("### 📚 Knowledge Base")
    
    openai_api_key = st.text_input(
        "🔑 OpenAI API Key (Optional)", 
        type="password", 
        help="Enter your OpenAI API key for advanced AI features. Leave empty for demo mode."
    )
    
    if openai_api_key:
        st.session_state.demo_mode = False
        st.success("✅ AI Mode Active")
    else:
        st.session_state.demo_mode = True
        st.info("ℹ️ Demo Mode Active")
    
    st.markdown("---")
    
    # Configuration section
    with st.expander("⚙️ System Configuration"):
        st.markdown("**📧 Email Settings**")
        gmail_email = st.text_input("Gmail Address", value=st.session_state.gmail_config['email'])
        
        st.markdown("**💬 WhatsApp Settings**")
        whatsapp_phone = st.text_input("WhatsApp Number", value=st.session_state.whatsapp_config['phone_number'])
        
        if st.button("💾 Save Configuration"):
            st.session_state.gmail_config['email'] = gmail_email
            st.session_state.whatsapp_config['phone_number'] = whatsapp_phone
            st.session_state.whatsapp_config['wa_link'] = f'https://wa.me/{whatsapp_phone.replace("+", "")}'
            st.success("✅ Configuration saved!")
    
    st.markdown("---")
    
    # File upload section
    st.markdown("**📁 Upload Knowledge Base**")
    uploaded_files = st.file_uploader(
        "Upload documents",
        type=['pdf', 'docx', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'],
        accept_multiple_files=True,
        help="Supported: PDF, DOCX, Excel, Images"
    )
    
    # Process button
    if st.button("🚀 Process Knowledge Base", type="primary", use_container_width=True):
        with st.spinner("🔄 Processing knowledge base..."):
            all_text = ""
            faq_data = []
            
            # Process uploaded files
            if uploaded_files:
                progress_bar = st.progress(0)
                for idx, file in enumerate(uploaded_files):
                    file_type = file.name.split('.')[-1].lower()
                    st.info(f"📄 Processing: {file.name}")
                    
                    if file_type == 'pdf':
                        all_text += f"\n[PDF: {file.name}]\n{extract_text_from_pdf(file)}\n\n"
                    elif file_type == 'docx':
                        all_text += f"\n[DOCX: {file.name}]\n{extract_text_from_docx(file)}\n\n"
                    elif file_type in ['xlsx', 'xls']:
                        # Try to extract as FAQ first
                        faq_list = extract_faq_from_excel(file)
                        if faq_list:
                            faq_data.extend(faq_list)
                            st.success(f"✅ Loaded {len(faq_list)} FAQs from {file.name}")
                        
                        # Also extract as text
                        excel_text = extract_text_from_excel(file)
                        if excel_text:
                            all_text += f"\n[Excel: {file.name}]\n{excel_text}\n\n"
                    
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                progress_bar.empty()
            
            # Save FAQ data
            if faq_data:
                st.session_state.faq_database = faq_data
                st.success(f"✅ Loaded {len(faq_data)} FAQ entries!")
            
            # Create vector store if AI mode
            if all_text.strip() and not st.session_state.demo_mode and AI_SUPPORT:
                st.session_state.knowledge_base_text = all_text
                
                try:
                    from langchain.text_splitter import RecursiveCharacterTextSplitter
                    from langchain_community.vectorstores import FAISS
                    from langchain_openai import OpenAIEmbeddings
                    from langchain.docstore.document import Document as LangchainDocument
                    
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=1000,
                        chunk_overlap=200,
                        length_function=len
                    )
                    chunks = text_splitter.split_text(all_text)
                    
                    documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
                    
                    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
                    st.session_state.embeddings_model = embeddings
                    st.session_state.vector_store = FAISS.from_documents(documents, embeddings)
                    
                    st.success(f"✅ AI Vector Store created with {len(all_text):,} characters!")
                    st.balloons()
                except Exception as e:
                    st.error(f"❌ Vector store creation failed: {str(e)}")
            elif all_text.strip() or faq_data:
                st.session_state.kb_processed = True
                st.success("✅ Knowledge Base processed in Demo Mode!")
                st.balloons()
            else:
                st.warning("⚠️ No content to process. Please upload files.")
    
    # Knowledge base status
    st.markdown("---")
    st.markdown("**📊 Knowledge Base Status**")
    
    if st.session_state.vector_store or st.session_state.faq_database:
        st.success("✅ Active")
        if st.session_state.faq_database:
            st.metric("FAQs Loaded", len(st.session_state.faq_database))
        if st.session_state.vector_store:
            num_docs = st.session_state.vector_store.index.ntotal
            st.metric("Vector Docs", num_docs)
    else:
        st.info("ℹ️ Not loaded")
    
    # Export section
    st.markdown("---")
    st.markdown("**📥 Export Data**")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💬 Chats", use_container_width=True):
            if st.session_state.chat_history:
                df = pd.DataFrame(st.session_state.chat_history)
                csv = df.to_csv(index=False)
                st.download_button(
                    "⬇️ Download",
                    csv,
                    "chat_history.csv",
                    "text/csv",
                    use_container_width=True
                )
    
    with col2:
        if st.button("🎫 Tickets", use_container_width=True):
            if st.session_state.tickets:
                df = pd.DataFrame(st.session_state.tickets)
                csv = df.to_csv(index=False)
                st.download_button(
                    "⬇️ Download",
                    csv,
                    "tickets.csv",
                    "text/csv",
                    use_container_width=True
                )

# Main page header
st.markdown('<div class="main-header">🤖 AI Customer Support Agent Pro</div>', unsafe_allow_html=True)

# Demo mode banner
if st.session_state.demo_mode:
    st.markdown("""
    <div class="demo-mode-banner">
        ⚡ DEMO MODE ACTIVE | For full AI features, add OpenAI API Key in sidebar
    </div>
    """, unsafe_allow_html=True)

# Contact banner with JavaScript for cross-platform email support
st.markdown(f"""
<div class="contact-banner">
    <h3 style='color: white; text-align: center; margin-bottom: 20px; font-size: 1.8rem;'>
        📞 Get Support Instantly
    </h3>
    <div style='display: flex; justify-content: center; gap: 20px; flex-wrap: wrap;'>
        <!-- Email Support Card -->
        <div class="contact-card" onclick="openEmail()" style="cursor: pointer;">
            <h4>📧 Email Support</h4>
            <p>{st.session_state.gmail_config['email']}</p>
            <small style='color: #6C757D;'>24/7 Response</small>
            <div style='margin-top: 10px;'>
                <span class="email-button" style="font-size: 0.9rem; padding: 8px 16px;">✉️ Send Email</span>
            </div>
        </div>
        
        <!-- WhatsApp Support Card -->
        <div class="contact-card" onclick="openWhatsApp()" style="cursor: pointer;">
            <h4>💬 WhatsApp Support</h4>
            <p>{st.session_state.whatsapp_config['phone_number']}</p>
            <small style='color: #6C757D;'>Instant Messaging</small>
            <div style='margin-top: 10px;'>
                <span class="whatsapp-button" style="font-size: 0.9rem; padding: 8px 16px;">💬 Open WhatsApp</span>
            </div>
        </div>
    </div>
</div>

<script>
function openEmail() {{
    const email = '{st.session_state.gmail_config['email']}';
    const subject = 'Support Request - Customer Inquiry';
    const body = 'Hello Support Team,%0D%0A%0D%0AI need assistance with:%0D%0A%0D%0A[Please describe your issue here]%0D%0A%0D%0AThank you!';
    
    // Try multiple methods for maximum compatibility
    
    // Method 1: Try Gmail web (works on all platforms)
    const gmailUrl = `https://mail.google.com/mail/?view=cm&fs=1&to=${{email}}&su=${{encodeURIComponent(subject)}}&body=${{body}}`;
    
    // Method 2: Mailto fallback
    const mailtoUrl = `mailto:${{email}}?subject=${{encodeURIComponent(subject)}}&body=${{body}}`;
    
    // Check if on mobile or desktop
    const isMobile = /iPhone|iPad|iPod|Android/i.test(navigator.userAgent);
    
    if (isMobile) {{
        // On mobile: Try Gmail app first, then mailto
        window.location.href = mailtoUrl;
    }} else {{
        // On desktop: Open Gmail in new tab
        const newWindow = window.open(gmailUrl, '_blank');
        
        // If popup blocked, try mailto
        if (!newWindow || newWindow.closed || typeof newWindow.closed == 'undefined') {{
            window.location.href = mailtoUrl;
        }}
    }}
}}

function openWhatsApp() {{
    const phone = '{st.session_state.whatsapp_config['phone_number'].replace('+', '').replace('-', '').replace(' ', '')}';
    const message = 'Hello! I need support.';
    const whatsappUrl = `https://wa.me/${{phone}}?text=${{encodeURIComponent(message)}}`;
    
    window.open(whatsappUrl, '_blank');
}}
</script>
""", unsafe_allow_html=True)

# Channel metrics
st.markdown("### 📊 Channel Overview")
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.metric("🌐 Website", st.session_state.channel_stats.get('Website', 0))
with col2:
    st.metric("💬 WhatsApp", st.session_state.channel_stats.get('WhatsApp', 0))
with col3:
    st.metric("📧 Email", st.session_state.channel_stats.get('Email', 0))
with col4:
    st.metric("👤 Manual", st.session_state.channel_stats.get('Manual', 0))

st.markdown("---")

# Tabs
tab1, tab2, tab3, tab4 = st.tabs([
    "💬 Live Chat", 
    "📱 WhatsApp Support", 
    "🎫 Ticket System", 
    "📊 Analytics"
])

# Tab 1: Website Chat Agent
with tab1:
    st.markdown('<div class="sub-header">💬 Live Chat Interface</div>', unsafe_allow_html=True)
    st.caption("🌐 Real-time customer support chat")
    
    # Always show chat interface
    chat_container = st.container()
    
    with chat_container:
        for i, chat in enumerate(st.session_state.chat_history):
            if chat['role'] == 'user':
                st.markdown(
                    f'<div class="chat-message user-message">'
                    f'<strong>👤 Customer:</strong><br>{chat["message"]}'
                    f'<br><small style="color: #6C757D;">🌍 {chat.get("language", "English")} | '
                    f'📂 {chat.get("category", "General")}</small>'
                    f'</div>', 
                    unsafe_allow_html=True
                )
            else:
                confidence = chat.get('confidence', 0)
                response_time = chat.get('response_time', 0)
                confidence_badge = get_confidence_badge(confidence)
                
                st.markdown(
                    f'<div class="chat-message bot-message">'
                    f'<strong>🤖 AI Agent:</strong><br>{chat["message"]}<br><br>'
                    f'<div style="display: flex; gap: 10px; flex-wrap: wrap; margin-top: 10px;">'
                    f'{confidence_badge}'
                    f'<span style="color: #6C757D;">⏱️ {response_time:.2f}s</span>'
                    f'</div>'
                    f'</div>', 
                    unsafe_allow_html=True
                )
                
                # Feedback buttons
                if i not in st.session_state.feedback:
                    col1, col2, col3 = st.columns([1, 1, 10])
                    with col1:
                        if st.button("👍 Helpful", key=f"up_{i}"):
                            st.session_state.feedback[i] = 'positive'
                            st.rerun()
                    with col2:
                        if st.button("👎 Not Helpful", key=f"down_{i}"):
                            st.session_state.feedback[i] = 'negative'
                            st.rerun()
                else:
                    if st.session_state.feedback[i] == 'positive':
                        st.markdown('<div class="success-box">✓ Thank you for your feedback!</div>', unsafe_allow_html=True)
                    else:
                        st.markdown('<div class="warning-box">✗ We\'ll improve our responses</div>', unsafe_allow_html=True)
    
    # Input form
    st.markdown("---")
    with st.form(key="chat_form", clear_on_submit=True):
        col1, col2 = st.columns([5, 1])
        with col1:
            user_input = st.text_input(
                "Your message:", 
                placeholder="Type your question here... (Hindi/English supported)",
                label_visibility="collapsed"
            )
        with col2:
            submit_button = st.form_submit_button("Send 🚀", use_container_width=True, type="primary")
    
    if submit_button and user_input:
        start_time = time.time()
        
        # Detect language and category
        language = detect_language(user_input)
        category = categorize_query(user_input)
        
        # Update analytics
        st.session_state.analytics['total_queries'] += 1
        st.session_state.analytics['languages'][language] = st.session_state.analytics['languages'].get(language, 0) + 1
        st.session_state.analytics['categories'][category] += 1
        st.session_state.channel_stats['Website'] += 1
        
        # Save user message
        user_chat = {
            'role': 'user',
            'message': user_input,
            'timestamp': datetime.now(),
            'language': language,
            'category': category
        }
        st.session_state.chat_history.append(user_chat)
        save_chat_to_db(user_chat)
        
        # Get AI response
        with st.spinner(f"🤔 Thinking... (Language: {language})"):
            if st.session_state.demo_mode:
                # Demo mode - use FAQ matching
                answer, confidence = get_demo_response(user_input, st.session_state.faq_database)
            else:
                # AI mode
                answer, confidence, source_docs = get_ai_response(
                    user_input, 
                    st.session_state.vector_store, 
                    openai_api_key,
                    language
                )
        
        response_time = time.time() - start_time
        
        # Check if escalation needed
        if confidence < 0.6:
            ticket_id = create_ticket(user_input, language, category, confidence, 'Website')
            priority = assign_priority(confidence, category)
            escalation_msg = f"\n\n⚠️ **Escalation Notice:** Your query has been forwarded to our support team.\n\n📋 **Ticket Details:**\n- Ticket ID: `{ticket_id}`\n- Priority: {priority}\n- Expected Response: Within 24 hours"
            if language != 'English' and TRANSLATION_SUPPORT:
                escalation_msg = translate_text(escalation_msg, language)
            answer += escalation_msg
        else:
            st.session_state.analytics['answered'] += 1
        
        # Save bot response
        bot_chat = {
            'role': 'bot',
            'message': answer,
            'timestamp': datetime.now(),
            'confidence': confidence,
            'response_time': response_time,
            'language': language,
            'category': category
        }
        st.session_state.chat_history.append(bot_chat)
        save_chat_to_db(bot_chat)
        
        st.rerun()

# Tab 2: WhatsApp Support
with tab2:
    st.markdown('<div class="sub-header">📱 WhatsApp Integration</div>', unsafe_allow_html=True)
    st.caption(f"Connected: {st.session_state.whatsapp_config['phone_number']}")
    
    st.markdown("---")
    
    # Quick chat link
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("### 🚀 Quick Chat Link")
        quick_message = st.text_input(
            "Pre-fill message (optional)", 
            placeholder="Hello, I need assistance..."
        )
    with col2:
        st.write("")
        st.write("")
        wa_link = create_whatsapp_link(
            st.session_state.whatsapp_config['phone_number'], 
            quick_message if quick_message else "Hello! I need support."
        )
        st.markdown(
            f'<a href="{wa_link}" target="_blank" class="whatsapp-button">💬 Open WhatsApp</a>', 
            unsafe_allow_html=True
        )
    
    st.markdown("---")
    st.markdown("### 📊 WhatsApp Stats")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Messages", len(st.session_state.whatsapp_messages))
    with col2:
        wa_escalated = sum(1 for msg in st.session_state.whatsapp_messages if msg.get('ticket_id'))
        st.metric("Escalated", wa_escalated)
    with col3:
        wa_resolved = len(st.session_state.whatsapp_messages) - wa_escalated
        st.metric("Resolved", wa_resolved)

# Tab 3: Ticket System
with tab3:
    st.markdown('<div class="sub-header">🎫 Support Ticket System</div>', unsafe_allow_html=True)
    
    if st.session_state.tickets:
        # Filters
        col1, col2, col3 = st.columns(3)
        with col1:
            status_filter = st.multiselect(
                "Filter by Status",
                options=['Open', 'In Progress', 'Closed'],
                default=['Open', 'In Progress']
            )
        with col2:
            priority_filter = st.multiselect(
                "Filter by Priority",
                options=['High', 'Medium', 'Low'],
                default=['High', 'Medium', 'Low']
            )
        with col3:
            category_filter = st.multiselect(
                "Filter by Category",
                options=['Billing', 'Technical', 'General', 'Product'],
                default=['Billing', 'Technical', 'General', 'Product']
            )
        
        df_tickets = pd.DataFrame(st.session_state.tickets)
        
        filtered_df = df_tickets[
            (df_tickets['status'].isin(status_filter)) &
            (df_tickets['priority'].isin(priority_filter)) &
            (df_tickets['category'].isin(category_filter))
        ]
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Tickets", len(df_tickets))
        with col2:
            st.metric("Open", len(df_tickets[df_tickets['status'] == 'Open']))
        with col3:
            st.metric("High Priority", len(df_tickets[df_tickets['priority'] == 'High']))
        with col4:
            avg_resolution = df_tickets[df_tickets['resolution_time'].notna()]['resolution_time'].mean()
            st.metric("Avg Resolution", f"{avg_resolution:.1f}h" if not pd.isna(avg_resolution) else "N/A")
        
        st.markdown("---")
        
        # Ticket cards
        for idx, ticket in filtered_df.iterrows():
            priority_badge = get_priority_badge(ticket['priority'])
            status_badge = get_status_badge(ticket['status'])
            
            with st.expander(
                f"Ticket #{ticket['id']} | {ticket['category']} | {ticket['channel']}", 
                expanded=False
            ):
                st.markdown(f"""
                <div style='margin-bottom: 1rem;'>
                    {priority_badge} {status_badge}
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**Query:** {ticket['query']}")
                    st.markdown(f"**Category:** {ticket['category']}")
                    st.markdown(f"**Language:** {ticket['language']}")
                    st.markdown(f"**Channel:** {ticket['channel']}")
                with col2:
                    st.markdown(f"**Assigned To:** {ticket['assigned_to']}")
                    st.markdown(f"**Created:** {ticket['timestamp']}")
                    if ticket['resolved_at']:
                        st.markdown(f"**Resolved:** {ticket['resolved_at']}")
                        st.markdown(f"**Resolution Time:** {ticket['resolution_time']:.1f}h")
                
                st.markdown("---")
                
                col1, col2, col3 = st.columns([2, 2, 1])
                with col1:
                    new_status = st.selectbox(
                        "Update Status",
                        options=['Open', 'In Progress', 'Closed'],
                        index=['Open', 'In Progress', 'Closed'].index(ticket['status']),
                        key=f"status_{ticket['id']}"
                    )
                with col2:
                    new_assignee = st.selectbox(
                        "Reassign To",
                        options=['Agent A', 'Agent B', 'Agent C', 'Agent D'],
                        index=['Agent A', 'Agent B', 'Agent C', 'Agent D'].index(ticket['assigned_to']),
                        key=f"assign_{ticket['id']}"
                    )
                with col3:
                    st.write("")
                    st.write("")
                    if st.button("💾 Update", key=f"update_{ticket['id']}", type="primary"):
                        if new_status == 'Closed' and ticket['status'] != 'Closed':
                            created = datetime.strptime(ticket['timestamp'], "%Y-%m-%d %H:%M:%S")
                            resolution_time = (datetime.now() - created).total_seconds() / 3600
                            st.session_state.tickets[idx]['resolved_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            st.session_state.tickets[idx]['resolution_time'] = resolution_time
                        
                        st.session_state.tickets[idx]['status'] = new_status
                        st.session_state.tickets[idx]['assigned_to'] = new_assignee
                        
                        st.markdown('<div class="success-box">✅ Ticket updated!</div>', unsafe_allow_html=True)
                        time.sleep(1)
                        st.rerun()
    else:
        st.markdown('<div class="info-box">ℹ️ No tickets created yet. Tickets are auto-generated when confidence is low.</div>', unsafe_allow_html=True)

# Tab 4: Analytics
with tab4:
    st.markdown('<div class="sub-header">📊 Performance Analytics</div>', unsafe_allow_html=True)
    
    analytics = st.session_state.analytics
    
    # Key metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Queries", analytics.get('total_queries', 0))
    with col2:
        st.metric("Answered", analytics.get('answered', 0))
    with col3:
        st.metric("Escalated", analytics.get('escalated', 0))
    with col4:
        total = analytics.get('total_queries', 0)
        answered = analytics.get('answered', 0)
        resolution_rate = (answered / total * 100) if total > 0 else 0
        st.metric("Resolution Rate", f"{resolution_rate:.1f}%")
    
    st.markdown("---")
    
    # Charts
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📊 Query Distribution")
        answered = analytics.get('answered', 0)
        escalated = analytics.get('escalated', 0)
        
        if answered > 0 or escalated > 0:
            fig_queries = go.Figure(data=[
                go.Bar(
                    x=['Auto-Resolved', 'Escalated'],
                    y=[answered, escalated],
                    marker_color=['#00A86B', '#DC143C'],
                    text=[answered, escalated],
                    textposition='auto'
                )
            ])
            fig_queries.update_layout(
                height=350,
                plot_bgcolor='white'
            )
            st.plotly_chart(fig_queries, use_container_width=True)
        else:
            st.info("No data yet")
    
    with col2:
        st.markdown("### 📂 Category Breakdown")
        categories = analytics.get('categories', {})
        category_df = pd.DataFrame({
            'Category': list(categories.keys()),
            'Count': list(categories.values())
        })
        
        if category_df['Count'].sum() > 0:
            fig_category = px.bar(
                category_df,
                x='Category',
                y='Count',
                color='Category'
            )
            fig_category.update_layout(height=350, showlegend=False)
            st.plotly_chart(fig_category, use_container_width=True)
        else:
            st.info("No data yet")

# Footer
st.markdown("---")
st.markdown(f"""
<div class="footer">
    <h3 style='color: var(--primary-dark); margin-bottom: 1rem;'>🤖 AI Customer Support Agent Pro v5.0</h3>
    <p style='font-weight: 600; margin-bottom: 0.5rem;'>Powered by OpenAI GPT & Streamlit</p>
    <p style='margin-bottom: 1rem;'>
        📧 {st.session_state.gmail_config['email']} | 
        💬 {st.session_state.whatsapp_config['phone_number']}
    </p>
    <div style='background: var(--bg-light); padding: 1rem; border-radius: 8px; margin-top: 1rem;'>
        <p style='margin: 0; font-size: 0.9rem;'>
            <strong>✅ Features:</strong> Multi-Channel Support | AI & Demo Mode | 
            Document Processing | Multilingual | Smart Tickets | Real-time Analytics
        </p>
    </div>
</div>
""", unsafe_allow_html=True)
